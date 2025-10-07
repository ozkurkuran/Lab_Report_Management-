"""
Reports API - DOCX/PDF/XLSX üretimi
"""
import io
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Optional
from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import FileResponse, StreamingResponse
from sqlmodel import Session, select
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.database import get_session, DATA_DIR
from app.models import Entry, Experiment, Project, Attachment, Dataset, Chart, Template
from app.schemas import ReportDOCXRequest, ReportPDFRequest, ReportXLSXRequest

router = APIRouter()

# Report dizini
REPORT_DIR = DATA_DIR / "storage" / "reports"
REPORT_DIR.mkdir(parents=True, exist_ok=True)


@router.post("/docx")
async def generate_docx_report(
    request: ReportDOCXRequest,
    session: Session = Depends(get_session),
):
    """DOCX raporu oluştur"""
    
    # Entry ve ilişkili verileri getir
    entry = session.get(Entry, request.entry_id)
    if not entry:
        raise HTTPException(status_code=404, detail="Entry not found")
    
    experiment = session.get(Experiment, entry.experiment_id)
    project = session.get(Project, experiment.project_id) if experiment else None
    
    # DOCX oluştur
    doc = Document()
    
    # Başlık
    title = doc.add_heading(entry.title, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata tablosu
    doc.add_heading('Rapor Bilgileri', level=1)
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Light Grid Accent 1'
    
    table.cell(0, 0).text = 'Proje'
    table.cell(0, 1).text = project.name if project else 'N/A'
    
    table.cell(1, 0).text = 'Deney'
    table.cell(1, 1).text = experiment.title if experiment else 'N/A'
    
    table.cell(2, 0).text = 'Entry'
    table.cell(2, 1).text = entry.title
    
    table.cell(3, 0).text = 'Tarih'
    table.cell(3, 1).text = entry.created_at.strftime('%Y-%m-%d %H:%M')
    
    table.cell(4, 0).text = 'Versiyon'
    table.cell(4, 1).text = str(entry.version)
    
    table.cell(5, 0).text = 'Etiketler'
    table.cell(5, 1).text = ', '.join(entry.tags) if entry.tags else '-'
    
    # İçerik
    doc.add_heading('İçerik', level=1)
    
    # Markdown içeriği paragraf olarak ekle (basit parsing)
    lines = entry.body_md.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        if line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('| '):
            # Basit tablo desteği (TODO: geliştirilecek)
            doc.add_paragraph(line)
        else:
            doc.add_paragraph(line)
    
    # Ekleri listele
    attachments = session.exec(
        select(Attachment).where(Attachment.entry_id == request.entry_id)
    ).all()
    
    if attachments:
        doc.add_page_break()
        doc.add_heading('Ekler', level=1)
        
        for i, att in enumerate(attachments, 1):
            doc.add_heading(f'Ek {i}: {att.original_name}', level=2)
            if att.caption:
                doc.add_paragraph(att.caption)
            
            # Resim ekleri gömülü olarak ekle
            if att.file_type in ['png', 'jpg', 'jpeg']:
                try:
                    img_path = DATA_DIR / att.file_path
                    if img_path.exists():
                        doc.add_picture(str(img_path), width=Inches(5))
                except Exception as e:
                    doc.add_paragraph(f'[Resim yüklenemedi: {str(e)}]')
    
    # Dataset ve grafikleri ekle
    datasets = session.exec(
        select(Dataset).where(Dataset.entry_id == request.entry_id)
    ).all()
    
    if datasets:
        doc.add_page_break()
        doc.add_heading('Veri Setleri ve Grafikler', level=1)
        
        for dataset in datasets:
            doc.add_heading(f'Dataset: {dataset.name}', level=2)
            doc.add_paragraph(f'Satır sayısı: {dataset.row_count}')
            
            # İstatistikler
            if dataset.stats_json:
                doc.add_paragraph('İstatistikler:', style='Intense Quote')
                for col, stats in dataset.stats_json.items():
                    doc.add_paragraph(
                        f"  {col}: mean={stats.get('mean', 'N/A'):.2f}, "
                        f"std={stats.get('std', 'N/A'):.2f}",
                        style='List Bullet'
                    )
            
            # Grafikleri ekle
            charts = session.exec(
                select(Chart).where(Chart.dataset_id == dataset.id)
            ).all()
            
            for chart in charts:
                try:
                    chart_path = DATA_DIR / chart.image_path
                    if chart_path.exists():
                        doc.add_picture(str(chart_path), width=Inches(5.5))
                        if chart.title:
                            p = doc.add_paragraph(chart.title)
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    doc.add_paragraph(f'[Grafik yüklenemedi: {str(e)}]')
    
    # Kaydet
    timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    filename = f"report_entry_{request.entry_id}_{timestamp}.docx"
    file_path = REPORT_DIR / filename
    
    doc.save(str(file_path))
    
    return FileResponse(
        path=file_path,
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


@router.post("/xlsx")
async def generate_xlsx_report(
    request: ReportXLSXRequest,
    session: Session = Depends(get_session),
):
    """XLSX raporu oluştur (dataset verilerini içerir)"""
    
    entry = session.get(Entry, request.entry_id)
    if not entry:
        raise HTTPException(status_code=404, detail="Entry not found")
    
    experiment = session.get(Experiment, entry.experiment_id)
    project = session.get(Project, experiment.project_id) if experiment else None
    
    # Excel dosyası oluştur
    timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    filename = f"report_entry_{request.entry_id}_{timestamp}.xlsx"
    file_path = REPORT_DIR / filename
    
    with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
        # Özet sayfası
        summary_data = {
            'Alan': ['Proje', 'Deney', 'Entry', 'Tarih', 'Versiyon', 'Etiketler'],
            'Değer': [
                project.name if project else 'N/A',
                experiment.title if experiment else 'N/A',
                entry.title,
                entry.created_at.strftime('%Y-%m-%d %H:%M'),
                str(entry.version),
                ', '.join(entry.tags) if entry.tags else '-'
            ]
        }
        summary_df = pd.DataFrame(summary_data)
        summary_df.to_excel(writer, sheet_name='Özet', index=False)
        
        # Dataset'leri ekle
        datasets = session.exec(
            select(Dataset).where(Dataset.entry_id == request.entry_id)
        ).all()
        
        for i, dataset in enumerate(datasets, 1):
            sheet_name = f'Dataset_{i}'[:31]  # Excel limit: 31 char
            
            try:
                file_path_dataset = DATA_DIR / dataset.source_file
                if dataset.source_file.endswith('.csv'):
                    df = pd.read_csv(file_path_dataset)
                else:
                    df = pd.read_excel(file_path_dataset)
                
                df.to_excel(writer, sheet_name=sheet_name, index=False)
            except Exception as e:
                error_df = pd.DataFrame({'Hata': [f'Dataset yüklenemedi: {str(e)}']})
                error_df.to_excel(writer, sheet_name=sheet_name, index=False)
    
    return FileResponse(
        path=file_path,
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )


@router.post("/pdf")
async def generate_pdf_report(
    request: ReportPDFRequest,
    session: Session = Depends(get_session),
):
    """PDF raporu oluştur (basit HTML → PDF)"""
    
    entry = session.get(Entry, request.entry_id)
    if not entry:
        raise HTTPException(status_code=404, detail="Entry not found")
    
    experiment = session.get(Experiment, entry.experiment_id)
    project = session.get(Project, experiment.project_id) if experiment else None
    
    # Basit HTML oluştur
    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="utf-8">
        <style>
            body {{ font-family: Arial, sans-serif; margin: 40px; }}
            h1 {{ color: #2c3e50; border-bottom: 2px solid #3498db; }}
            h2 {{ color: #34495e; margin-top: 30px; }}
            table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
            th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
            th {{ background-color: #3498db; color: white; }}
            .metadata {{ background-color: #f8f9fa; padding: 15px; border-radius: 5px; }}
        </style>
    </head>
    <body>
        <h1>{entry.title}</h1>
        
        <div class="metadata">
            <table>
                <tr><th>Proje</th><td>{project.name if project else 'N/A'}</td></tr>
                <tr><th>Deney</th><td>{experiment.title if experiment else 'N/A'}</td></tr>
                <tr><th>Tarih</th><td>{entry.created_at.strftime('%Y-%m-%d %H:%M')}</td></tr>
                <tr><th>Versiyon</th><td>{entry.version}</td></tr>
                <tr><th>Etiketler</th><td>{', '.join(entry.tags) if entry.tags else '-'}</td></tr>
            </table>
        </div>
        
        <h2>İçerik</h2>
        <pre style="white-space: pre-wrap;">{entry.body_md}</pre>
    </body>
    </html>
    """
    
    # HTML'i PDF'e çevir (basit versiyon - WeasyPrint gerektirmeden)
    timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    filename = f"report_entry_{request.entry_id}_{timestamp}.html"
    file_path = REPORT_DIR / filename
    
    with open(file_path, 'w', encoding='utf-8') as f:
        f.write(html_content)
    
    return FileResponse(
        path=file_path,
        filename=filename,
        media_type="text/html"
    )


@router.get("/export/experiment/{experiment_id}/zip")
async def export_experiment_zip(
    experiment_id: int,
    session: Session = Depends(get_session),
):
    """Deneyi tüm ekleriyle ZIP olarak indir"""
    
    experiment = session.get(Experiment, experiment_id)
    if not experiment:
        raise HTTPException(status_code=404, detail="Experiment not found")
    
    # Entry'leri ve eklerini topla
    entries = session.exec(
        select(Entry).where(Entry.experiment_id == experiment_id)
    ).all()
    
    # Bellek içi ZIP oluştur
    zip_buffer = io.BytesIO()
    
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        # Her entry için
        for entry in entries:
            entry_dir = f"entry_{entry.id}_{entry.title[:20]}"
            
            # Entry içeriğini txt olarak ekle
            entry_content = f"""
Entry: {entry.title}
Tarih: {entry.created_at}
Versiyon: {entry.version}
Etiketler: {', '.join(entry.tags)}

{entry.body_md}
"""
            zip_file.writestr(f"{entry_dir}/entry.txt", entry_content)
            
            # Attachments
            attachments = session.exec(
                select(Attachment).where(Attachment.entry_id == entry.id)
            ).all()
            
            for att in attachments:
                try:
                    att_path = DATA_DIR / att.file_path
                    if att_path.exists():
                        zip_file.write(
                            att_path,
                            f"{entry_dir}/attachments/{att.original_name}"
                        )
                except Exception:
                    pass
            
            # Datasets
            datasets = session.exec(
                select(Dataset).where(Dataset.entry_id == entry.id)
            ).all()
            
            for dataset in datasets:
                try:
                    ds_path = DATA_DIR / dataset.source_file
                    if ds_path.exists():
                        zip_file.write(
                            ds_path,
                            f"{entry_dir}/datasets/{dataset.name}{ds_path.suffix}"
                        )
                except Exception:
                    pass
    
    zip_buffer.seek(0)
    
    filename = f"experiment_{experiment_id}_{experiment.title[:20]}.zip"
    
    return StreamingResponse(
        zip_buffer,
        media_type="application/zip",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )
